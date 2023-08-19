from docx import Document
from docx.oxml import OxmlElement
from docx.enum.text import  WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
def remove_table_shading_and_add_borders(table):
    tblPr = table._element.tblPr

    # Remove shading from table
    tblShading = tblPr.find(qn('w:shd'))
    if tblShading is not None:
        tblPr.remove(tblShading)

    for row in table.rows:
        for cell in row.cells:
            remove_cell_shading_and_add_border(cell)

def remove_cell_shading_and_add_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()

    # Remove shading
    shading = tcPr.find(qn('w:shd'))
    if shading is not None:
        tcPr.remove(shading)

    # Add borders
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for border_type in ['left', 'top', 'right', 'bottom']:
        border = tcBorders.find(qn(f'w:{border_type}'))
        if border is None:
            border = OxmlElement(f'w:{border_type}')
            tcBorders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')  # Black color

def process_docx_file(file_path):
    doc = Document(file_path)
    table = doc.tables[0]
    table.add_row()

    for i,cells in enumerate(table.rows[-1].cells):
            if i == 0:
                cells.text = "疑似辍学原因"
            elif i == 1:
                cells.text = "无"
    insertion_row = table.rows[4]._tr
    insertion_row.addnext(table.rows[-1]._tr)
    table = doc.tables[0]
    a = table.cell(5, 1)
    b = table.cell(5, 7)
    a.merge(b)
    # 获取第2行第3列的单元格
    cell = table.cell(2, 2)  # 注意索引是从0开始
    # 拆分单元格为2个
    new_tc1 = OxmlElement('w:tc')
    new_tcPr1 = OxmlElement('w:tcPr')
    new_tc1.append(new_tcPr1)
    cell._tc.getparent().insert(cell._tc.getparent().index(cell._tc), new_tc1)
    new_tc1.append(OxmlElement('w:p'))
    new_tc1.width = cell.width / 2

    new_tc2 = OxmlElement('w:tc')
    new_tcPr2 = OxmlElement('w:tcPr')
    new_tc2.append(new_tcPr2)
    cell._tc.getparent().insert(cell._tc.getparent().index(cell._tc), new_tc2)
    new_tc2.append(OxmlElement('w:p'))
    new_tc2.width = cell.width / 2

    # 在第一个新单元格中写入文字
    new_tc1_paragraph = new_tc1.find(qn('w:p'))
    new_tc1_run = OxmlElement('w:r')
    new_tc1_text = OxmlElement('w:t')
    new_tc1_text.text = "是否为疑似辍学对象"
    new_tc1_run.append(new_tc1_text)
    new_tc1_paragraph.append(new_tc1_run)
    # 设置段落居中对齐
    new_tc1_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 在第二个新单元格中写入文字
    new_tc2_paragraph = new_tc2.find(qn('w:p'))
    new_tc2_run = OxmlElement('w:r')
    new_tc2_text = OxmlElement('w:t')
    new_tc2_text.text = "否"
    new_tc2_run.append(new_tc2_text)
    new_tc2_paragraph.append(new_tc2_run)
    # 设置段落居中对齐
    new_tc2_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 保存文档
    remove_table_shading_and_add_borders(table)
    doc.save(file_path)

def process_docx_files_in_folder(folder_path):
    for root, dirs, files in os.walk(folder_path):
        for filename in files:
            if filename.endswith('.docx'):
                file_path = os.path.join(root, filename)
                process_docx_file(file_path)

# 指定根文件夹路径
root_folder = '/Users/liangkaixi/Desktop/222/'

# 批量处理所有子文件夹中的 .docx 文件
process_docx_files_in_folder(root_folder)